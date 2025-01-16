import sys
import re
import os
import time
import mysql.connector
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QPushButton, QLineEdit, 
    QLabel, QVBoxLayout, QWidget, QInputDialog, 
    QMessageBox, QFileDialog, QDesktopWidget
)
from PyQt5.QtGui import QIcon
from cryptography.fernet import Fernet
import pandas as pd
from docx import Document
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, Alignment
from openpyxl.chart import BarChart, Reference
from pynput.mouse import Controller
import pyautogui
import keyboard


# ===================== USER AUTHENTICATION CLASS ====================



def setup_global_styles(app):
    app.setStyleSheet("""
        QInputDialog {
            background-color: black;
        }
        QInputDialog QLabel {
            color: white;
        }
        QInputDialog QLineEdit {
            background-color: #333;
            color: white;
            border: 1px solid #555;
        }
        QInputDialog QPushButton {
            background-color: #444;
            color: white;
            border: 1px solid #555;
            padding: 5px;
            min-width: 80px;
        }
        QInputDialog QPushButton:hover {
            background-color: #555;
        }
                       QMessageBox {
            background-color: rgb(42, 63, 84);
        }
        QMessageBox QPushButton {
            background-color: qlineargradient(
                spread:pad, x1:0, y1:0, x2:1, y2:1, 
                stop:0 #4682B4, stop:1 #1E3A5F
            );
            color: white;
            border: 2px solid #2A5F8E;
            border-radius: 10px;
            padding: 8px;
            min-width: 80px;
        }
        QMessageBox QPushButton:hover {
            background-color: qlineargradient(
                spread:pad, x1:0, y1:0, x2:1, y2:1, 
                stop:0 #3A6A99, stop:1 #144E75
            );
        }
        QMessageBox QLabel {
            color: white;
            font-size: 14px;
            font-weight: bold;
            font-family: Arial;
        }

    """)


class AutoClicker(QMainWindow):

    def __init__(self):
        super().__init__()

        self.setWindowTitle("Auto Clicker")
        self.mouse_controller = Controller()
        self.clicks = [] 
        
                
        self.single_click_button = QPushButton("Single Click", self)
        self.double_click_button = QPushButton("Double Click", self)
        self.right_click_button = QPushButton("Right Click",self)
        
        
        # Message label to show status or warnings
        self.message_label = QLabel("", self)
        self.message_label.setStyleSheet("color: red; font-weight: bold;")
        
        layout = QVBoxLayout()
        layout.addWidget(self.right_click_button)
        layout.addWidget(self.single_click_button)
        layout.addWidget(self.double_click_button)
        layout.addWidget(self.message_label)
                
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)
        
        self.set_stylesheet()  # apply the stylesheet

        self.single_click_button.clicked.connect(self.single_click)
        self.double_click_button.clicked.connect(self.double_click)
        self.right_click_button.clicked.connect(self.right_click)
        

    def set_stylesheet(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: #2F3C4F;
            }
            QPushButton {
                background-color: #4682B4;
                color: #FFFFFF;
                border: 1px solid #1E3A5F;
                border-radius: 5px;
                padding: 5px;
                font-size: 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #3A6A99;
            }
            

        """)

    def single_click(self):
        delay, ok = QInputDialog.getInt(self, "Single Click Delay", "Enter delay time in seconds:", 0, 0, 10)
        if ok:
            self.message_label.setText(f"Press q to stop single clicking... Delay: {delay} seconds")
            self.is_clicking = True
            while self.is_clicking and not keyboard.is_pressed('q'):
                pyautogui.click()  
                time.sleep(delay)  # for delay between clicks
            self.message_label.setText("Single clicking stopped!")

    def double_click(self):
        delay, ok = QInputDialog.getInt(self, "Double Click Delay", "Enter delay time in seconds:", 0, 0, 10)
        if ok:
            self.message_label.setText(f"Press q to stop double clicking... Delay: {delay} seconds")
            self.is_clicking = True
            while self.is_clicking and not keyboard.is_pressed('q'):
                pyautogui.doubleClick()
                time.sleep(delay)  
            self.message_label.setText("Double clicking stopped!")

    def right_click(self):
        delay, ok = QInputDialog.getInt(self, " Click Delay", "Enter delay time in seconds:", 0, 0, 10)
        if ok:
            self.message_label.setText(f"Press q to stop right clicking... Delay: {delay} seconds")
            self.is_clicking = True
            while self.is_clicking and not keyboard.is_pressed('q'):
                pyautogui.rightClick()
                time.sleep(delay)  
            self.message_label.setText("Right clicking stopped!")

    


def validate_password(password):
    # check if the password matches the conditions
        if len(password) < 8:
            return False, "Password must be at least 8 characters long."
        elif not re.search(r"[a-z]", password):  # at least one lowercase letter
            return False, "Password must contain at least one lowercase letter."
        elif not re.search(r"[A-Z]", password):  # at least one uppercase letter
            return False, "Password must contain at least one uppercase letter."
        elif not re.search(r"\d", password):  # at least one number
            return False, "Password must contain at least one number."
        elif not re.search(r"[!@#$%^&*(),.?\":{}|<>]", password):  # at least one symbol
            return False, "Password must contain at least one symbol."
        else:
            return True, "Password is valid!"

class UserAuth:
    def __init__(self):
        self.connection = mysql.connector.connect(
            host="localhost",
            user="root",
            password="",
            database="automation_app"
        )
        self.cursor = self.connection.cursor()
        
        # load or generate encryption key
        self.key_file = "key.key"
        self.key = self.load_or_generate_key()
        self.cipher_suite = Fernet(self.key)

    

    def load_or_generate_key(self):
        """Load the key from a file or generate a new one if it doesn't exist."""
        if os.path.exists(self.key_file):
            with open(self.key_file, "rb") as file:
                return file.read()
        else:
            key = Fernet.generate_key()
            with open(self.key_file, "wb") as file:
                file.write(key)
            return key

    def create_user_table(self):
        self.cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INT AUTO_INCREMENT PRIMARY KEY,
            username VARCHAR(255) UNIQUE,
            password VARCHAR(255)
        )
        """)
        self.connection.commit()

    def register_user(self, username, password):
        is_valid, password_message = validate_password(password)
        if not is_valid:
            return password_message

        # Check if username is empty or too short
        if not username:
            return "Username cannot be empty."
        if len(username) < 3:
            return "Username must be at least 3 characters long."

        # Check if the username already exists
        self.cursor.execute("SELECT * FROM users WHERE username = %s", (username,))
        if self.cursor.fetchone():
            return "Username already exists. Please choose a different username."

        # Encrypt the password
        encrypted_password = self.cipher_suite.encrypt(password.encode()).decode()

        try:
            # insert the new user into the database
            self.cursor.execute("INSERT INTO users (username, password) VALUES (%s, %s)", (username, encrypted_password))
            self.connection.commit()

            # check if the commit was successful
            if self.cursor.rowcount == 1:
                return "User registered successfully!"
            else:
                return "Error: User registration failed, no rows affected."
                
        except mysql.connector.Error as err:
            print(f"Error: {err}")
            return f"Database Error: {err}"

    def login_user(self, username, password):
        self.cursor.execute("SELECT password FROM users WHERE username = %s", (username,))
        result = self.cursor.fetchone()
        if result:
            try:
                decrypted_password = self.cipher_suite.decrypt(result[0].encode()).decode()
                return decrypted_password == password
            except Exception as e:
                print(f"Decryption failed: {e}")
                return False
        return False

    

# ===================== MAIN GUI CLASS ====================


class MainGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Automation Tool")
        self.setGeometry(200, 200, 400, 300)
        
        self.label = QLabel("Username:")
        self.username_input = QLineEdit()
        self.password_input = QLineEdit()
        self.password_input.setPlaceholderText("Password")
        self.login_button = QPushButton("Login")
        self.register_button = QPushButton("Register")
        
        self.message_label = QLabel("")
        self.message_label.setStyleSheet("color: red; font-weight: bold;")       


        layout = QVBoxLayout()
        layout.addWidget(self.label)
        layout.addWidget(self.username_input)
        layout.addWidget(self.password_input)
        layout.addWidget(self.login_button)
        layout.addWidget(self.register_button)
        layout.addWidget(self.message_label)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.set_stylesheet()  

        self.login_button.clicked.connect(self.login_user)
        self.register_button.clicked.connect(self.register_user)
        self.user_auth = UserAuth()
        self.user_auth.create_user_table()
        self.username_input.returnPressed.connect(self.login_user)
        self.password_input.returnPressed.connect(self.login_user)
        
    def set_stylesheet(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: qlineargradient(
                    spread:pad, x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #2A3F54, stop:1 #3B4D61
                );
                border: 1px solid #1E3A5F;
            }
            QLabel {
                color: #E1E8ED;
                font-size: 16px;
                font-weight: bold;
                font-family: Arial, Helvetica, sans-serif;
            }
            QLineEdit {
                background-color: #FFFFFF;
                border: 2px solid #5A9BD3;
                border-radius: 8px;
                padding: 6px;
                font-size: 14px;
                font-family: Arial, Helvetica, sans-serif;
                color:rgb(6, 74, 141);
            }
            QPushButton {
                background-color: qlineargradient(
                    spread:pad, x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #4682B4, stop:1 #1E3A5F
                );
                color: #FFFFFF;
                border: 2px solid #2A5F8E;
                border-radius: 10px;
                padding: 8px;
                font-size: 14px;
                font-weight: bold;
                font-family: Arial, Helvetica, sans-serif;
            }
            QPushButton:hover {
                background-color: qlineargradient(
                    spread:pad, x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #3A6A99, stop:1 #144E75
                );
                border: 2px solid #1A4C75;
            }
            QPushButton:pressed {
                background-color: #1A4C75;
            }
        """)

        


    def login_user(self):
        username = self.username_input.text()
        password = self.password_input.text()
        if self.user_auth.login_user(username, password):
            self.label.setText("Login Successful!")
            self.open_task_manager()
        else:
            self.label.setText("Wrong password or username!")

    def register_user(self):
        username = self.username_input.text()
        password = self.password_input.text()
        message = self.user_auth.register_user(username, password)

        if message == "User registered successfully!":
            self.message_label.setStyleSheet("color: green; font-weight: bold;")  
        else:
            self.message_label.setStyleSheet("color: red; font-weight: bold;")
        
        self.message_label.setText(message)


    def open_task_manager(self):
        from PyQt5.QtWidgets import QDesktopWidget
        self.close()
        self.set_stylesheet()
        self.task_window = TaskManager()
        self.task_window.show()
        # Center the Task Manager on screen
        qr = self.task_window.frameGeometry()
        cp = QDesktopWidget().availableGeometry().center()
        qr.moveCenter(cp)
        self.task_window.move(qr.topLeft())




# ===================== TASK MANAGER CLASS ====================



class TaskManager(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Task Manager")
        self.setGeometry(300, 300, 500, 400)
        
        self.excel_btn = QPushButton("Excel Automation", self)
        self.word_btn = QPushButton("Word Automation", self)
        self.click_btn = QPushButton("Auto Click", self)
        self.write_btn = QPushButton("Auto Write",self)
        
        self.excel_btn.clicked.connect(self.run_excel_automation)
        self.word_btn.clicked.connect(self.run_word_automation)
        self.click_btn.clicked.connect(self.run_auto_click)
        self.write_btn.clicked.connect(self.run_auto_write)

        icon_path = "icons/" 
        self.setWindowIcon(QIcon("icons/task_manager_icon.svg"))
        self.excel_btn.setIcon(QIcon(f"{icon_path}excel_icon.svg"))
        self.word_btn.setIcon(QIcon(f"{icon_path}word_icon.svg"))
        self.click_btn.setIcon(QIcon(f"{icon_path}click_icon.svg"))
        self.write_btn.setIcon(QIcon(f"{icon_path}write_icon.svg"))
        
        layout = QVBoxLayout()
        layout.addWidget(self.excel_btn)
        layout.addWidget(self.word_btn)
        layout.addWidget(self.click_btn)
        layout.addWidget(self.write_btn)
        
        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)  
        self.set_stylesheet()

    def set_stylesheet(self):
        self.setStyleSheet("""
            QMainWindow {
                background-color: qlineargradient(
                    spread:pad, x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #2A3F54, stop:1 #3B4D61
                );
                border: 1px solid #1E3A5F;
            }
            QLabel {
                color: #E1E8ED;
                font-size: 16px;
                font-weight: bold;
                font-family: Arial, Helvetica, sans-serif;
            }
            QLineEdit {
                background-color: #FFFFFF;
                border: 2px solid #5A9BD3;
                border-radius: 8px;
                padding: 6px;
                font-size: 14px;
                font-family: Arial, Helvetica, sans-serif;
                color: #2A3F54;
            }
            QPushButton {
                background-color: qlineargradient(
                    spread:pad, x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #4682B4, stop:1 #1E3A5F
                );
                color: #FFFFFF;
                border: 2px solid #2A5F8E;
                border-radius: 10px;
                padding: 8px;
                font-size: 14px;
                font-weight: bold;
                font-family: Arial, Helvetica, sans-serif;
            }
            QPushButton:hover {
                background-color: qlineargradient(
                    spread:pad, x1:0, y1:0, x2:1, y2:1, 
                    stop:0 #3A6A99, stop:1 #144E75
                );
                border: 2px solid #1A4C75;
            }
            QPushButton:pressed {
                background-color: #1A4C75;
            }
        """)

    def run_excel_automation(self):
        services = [
            "Import/Export Data", "Clean Data", "Convert Units", "Generate Visual Financial Report"
        ]
        
        service, ok = QInputDialog.getItem(self, "Select Excel Service", "Choose a service:", services, 0, False)
        
        if not ok:
            QMessageBox.warning(self, "Warning", "Operation cancelled!") 
            return

        if service == "Import/Export Data":
            self.import_export_data()
        elif service == "Clean Data":
            self.clean_data()
        elif service == "Convert Units":
            self.convert_units()
        elif service == "Generate Visual Financial Report":
                input_file, _ = QFileDialog.getOpenFileName(self, "Select Input Excel File", "", "Excel Files (*.xlsx *.xls)")
        
                if not input_file:
                    QMessageBox.warning(self, "Warning", "No input file selected. Operation cancelled!")
                    return  # Exit if no input file is selected

                try:
                    input_data = pd.read_excel(input_file)
                except Exception as e:
                    QMessageBox.critical(self, "Error", f"Failed to load the input file: {e}")
                    return  
                

                input_data.columns = input_data.columns.astype(str)

                output_file, _ = QFileDialog.getSaveFileName(self, "Save Financial Report", "", "Excel Files (*.xlsx)")
                
                if not output_file:
                    QMessageBox.warning(self, "Warning", "No output file selected. Operation cancelled!")
                    return  

                self.generate_visual_financial_report(input_data, output_file)
                
                

    def import_export_data(self):
        file_type, ok = QInputDialog.getItem(self, "Select File Type", "Choose the file type:", ["CSV", "TXT"], 0, False)
        if not ok:
            return

        input_file, _ = QFileDialog.getOpenFileName(self, "Select Input File", "", f"{file_type} Files (*.{file_type.lower()})")
        if not input_file:
            return

        output_file, _ = QFileDialog.getSaveFileName(self, "Save Output File", "", "Excel Files (*.xlsx)")
        if not output_file:
            return

        if file_type == "CSV":
            df = pd.read_csv(input_file, delimiter=r'\s+', engine='python')  
        elif file_type == "TXT":
            df = pd.read_table(input_file, sep=r'\s+', engine='python')  

        try:
            df.to_excel(output_file, index=False)
            QMessageBox.information(self, "Success", f"File exported as {output_file}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save the Excel file: {e}")


    
    def clean_data(self):
        input_file, _ = QFileDialog.getOpenFileName(self, "Select Excel File", "", "Excel Files (*.xlsx *.xls)")
        if not input_file:
            return  
        try:
            df = pd.read_excel(input_file, dtype=str) 
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to read the Excel file: {e}")
            return

        df_cleaned = df.apply(lambda x: x.str.strip() if x.dtype == "object" else x)  
        df_cleaned = df_cleaned.drop_duplicates()

        output_file, _ = QFileDialog.getSaveFileName(self, "Save Cleaned File", "", "Excel Files (*.xlsx)")
        if not output_file:
            return  

        try:
            df_cleaned.to_excel(output_file, index=False) 
            QMessageBox.information(self, "Success", f"Cleaned file saved as {output_file}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save the cleaned file: {e}")

    def convert_units(self):
        
        input_file, _ = QFileDialog.getOpenFileName(self, "Select Excel File", "", "Excel Files (*.xlsx *.xls)")
        if not input_file:
            return

        # read the Excel file
        try:
            df = pd.read_excel(input_file)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to read the Excel file: {e}")
            return

        # Check if the DataFrame has any columns to select
        if df.empty or len(df.columns) == 0:
            QMessageBox.warning(self, "Error", "No data available in the selected file.")
            return

        # Ask user for the column name that needs conversion
        column_name, ok = QInputDialog.getItem(self, "Select Column", "Choose the column to convert:", df.columns.tolist(), 0, False)
        if not ok or not column_name:
            return  # If user canceled or no column selected, exit

        # Ask user for the conversion factor (e.g., miles to kilometers etc.)
        conversion_factor, ok = QInputDialog.getDouble(self, "Conversion Factor", f"Enter the conversion factor for {column_name}:(like km to mile or km to meter or something like that)", 1.60934, 0, 100, 5)
        if not ok:
            return  # If user canceled or invalid factor, exit

        # Check if the selected column exists in the DataFrame
        if column_name not in df.columns:
            QMessageBox.warning(self, "Error", f"The column '{column_name}' does not exist in the data.")
            return

        # Convert units in the selected column
        try:
            df[column_name] = df[column_name] * conversion_factor
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error during conversion: {e}")
            return

        output_file, _ = QFileDialog.getSaveFileName(self, "Save Converted File", "", "Excel Files (*.xlsx)")
        if not output_file:
            return

        try:
            df.to_excel(output_file, index=False)
            QMessageBox.information(self, "Success", f"Converted file saved as {output_file}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to save the converted file: {e}")

    def generate_visual_financial_report(self, input_data, output_file):
            """
            Generates a visual financial report based on transactional data with a chart, 
            and appends it to the existing Excel file. Properly handles merged cells.
            
            :param input_data: A pandas DataFrame containing the data to visualize (from the input file).
            :param output_file: Path where the generated financial report will be saved (with added visualizations).
            """
            category_column, ok = QInputDialog.getItem(self, "Select Category Column", "Choose the column to categorize the data:",
                                                        input_data.columns.tolist(), 0, False)
            if not ok:
                return  

            amount_column, ok = QInputDialog.getItem(self, "Select Numerical Column", "Choose the column containing the numerical data:",
                                                        input_data.columns.tolist(), 0, False)
            if not ok:
                return  

            if category_column not in input_data.columns or amount_column not in input_data.columns:
                QMessageBox.critical(self, "Error", "The selected columns do not exist in the data.")
                return

            input_data = input_data[input_data[category_column] != category_column]
            
            input_data[category_column] = input_data[category_column].astype(str)  # Ensure it's a string
            input_data[amount_column] = pd.to_numeric(input_data[amount_column], errors='coerce')  # Ensure it's numeric

            input_data = input_data.dropna(subset=[amount_column])

            summary = input_data.groupby(category_column)[amount_column].sum().reset_index()

            try:
                wb = load_workbook(output_file)
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to open the Excel file: {e}")
                return

            if "Financial Report" in wb.sheetnames:
                wb.remove(wb["Financial Report"])
            
            ws = wb.create_sheet("Financial Report")

            title = "Automated Financial Visual Report"
            ws.merge_cells('A1:F1')
            ws['A1'] = title
            ws['A1'].font = Font(size=14, bold=True)
            ws['A1'].alignment = Alignment(horizontal="center")

            headers = [category_column, amount_column]
            ws.append(headers)

            for index, row in summary.iterrows():
                ws.append([row[category_column], row[amount_column]])

            chart = BarChart()
            data = Reference(ws, min_col=2, min_row=2, max_col=2, max_row=len(summary) + 2)
            categories = Reference(ws, min_col=1, min_row=2, max_row=len(summary) + 2)
            chart.add_data(data, titles_from_data=False)
            chart.set_categories(categories)
            chart.title = "Financial Breakdown"
            chart.x_axis.title = category_column
            chart.y_axis.title = amount_column

            ws.add_chart(chart, "E5")

            try:
                wb.save(output_file)
                QMessageBox.information(self, "Success", f"Report added to {output_file}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save the report: {e}")



    def run_auto_click(self):
        self.auto_clicker = AutoClicker()
        self.auto_clicker.show()
            

    def run_auto_write(self):
        txt_file, _ = QFileDialog.getOpenFileName(self, "Select Text File", "", "Text Files (*.txt)")
        if not txt_file:
            QMessageBox.warning(self, "Warning", "Operation Cancelled!")
            return
        
        try:
            with open(txt_file, "r", encoding="utf-8") as file:
                content = file.read()
                
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to read the file: {e}")
            return
        
        delay, ok = QInputDialog.getInt(self, "Typing Speed", "Enter delay between keystrokes in milliseconds:", 0, 0, 500)
        if not ok:
            QMessageBox.warning(self, "Warning", "Operation Cancelled!")
            return

        
        time.sleep(5)
        for char in content:
            keyboard.write(char) 
            time.sleep(delay / 1000.0)  
        
        QMessageBox.information(self, "Success", "Text Typed Successfully!")



    
    def run_word_automation(self):
        services = ["Auto Table Maker", "Text to Word", "CSV to Word"]
        service, ok = QInputDialog.getItem(self, "Select Word Service", "Choose a service:", services, 0, False)
        if not ok:
            QMessageBox.warning(self, "Warning", "Operation Cancelled!")
            return

        if service == "Auto Table Maker":
            self.auto_table_maker()
        elif service == "Text to Word":
            self.text_to_word()
        elif service == "CSV to Word":
            self.csv_to_word()

    def auto_table_maker(self):
        choice, ok = QInputDialog.getItem(
            self, "Table Input Option", 
            "Choose an input method for the table:", 
            ["Manual Input", "From Text File"], 0, False
        )
        if not ok or not choice:
            QMessageBox.warning(self, "Warning", "Operation Cancelled!")
            return
        
        doc = Document()
        
        if choice == "Manual Input":
            rows, ok1 = QInputDialog.getInt(self, "Input Rows", "Number of rows:", 3, 1, 100)
            cols, ok2 = QInputDialog.getInt(self, "Input Columns", "Number of columns:", 3, 1, 20)
            if not (ok1 and ok2):
                QMessageBox.warning(self, "Warning", "Operation Cancelled!")
                return
            
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            for i in range(rows):
                for j in range(cols):
                    table.cell(i, j).text = " "
                    
        elif choice == "From Text File":
            file_path, _ = QFileDialog.getOpenFileName(self, "Select Text File", "", "Text Files (*.txt)")
            if not file_path:
                QMessageBox.warning(self, "Warning", "Operation Cancelled!")
                return
            try:
                with open(file_path, "r", encoding="utf-8") as file:
                    lines = file.readlines()
                    
                rows_data = [line.strip().split() for line in lines if line.strip()]
                rows = len(rows_data)
                cols = max(len(row) for row in rows_data)
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                for i, row in enumerate(rows_data):
                    for j, value in enumerate(row):
                        table.cell(i, j).text = value
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to process the file: {e}")
                return
                
        
        file_name, ok = QInputDialog.getText(self, "Save As", "Enter the file name for the table:", QLineEdit.Normal, "auto_table")
        if ok and file_name:
            save_path= f"{file_name}.docx"
            doc.save(save_path)
            QMessageBox.information(self,"Success",f"Table saved as {save_path}")
        else:
            QMessageBox.warning(self,"Warnings","Operation Canceled! No file saved")
    def text_to_word(self):
        txt_file, _ = QFileDialog.getOpenFileName(self, "Select Text File", "", "Text Files (*.txt)")
        if not txt_file:
            return
        with open(txt_file, "r", encoding="utf-8") as f:
            content = f.read()
        doc = Document()
        doc.add_paragraph(content)
        doc.save("text_to_word.docx")
        QMessageBox.information(self, "Success", "File saved as 'text_to_word.docx'.")

    def csv_to_word(self):
        csv_file, _ = QFileDialog.getOpenFileName(self, "Select CSV File", "", "CSV Files (*.csv)")
        if not csv_file:
            return
        df = pd.read_csv(csv_file)
        doc = Document()
        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
        for j, col in enumerate(df.columns):
            table.cell(0, j).text = str(col)
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                table.cell(i+1, j).text = str(value)
        doc.save("csv_to_word.docx")
        QMessageBox.information(self, "Success", "File saved as 'csv_to_word.docx'.")


if __name__ == "__main__":
        
    app = QApplication(sys.argv)
    icon_path = "icons/" 
    app.setWindowIcon(QIcon("icons/task_manager_icon.svg"))

    setup_global_styles(app) 
    main_window = MainGUI()
    main_window.show()
    sys.exit(app.exec_())